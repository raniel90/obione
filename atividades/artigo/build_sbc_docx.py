#!/usr/bin/env python3
"""Gera o artigo ObiOne no template SBC (formato relatório/artigo).

Reproduz a formatação SBC: A4; margens 3,5 (topo) / 2,5 (base) / 3,0 (laterais) cm;
Times 12; título 16pt bold centralizado; autores 12pt bold centralizados; e-mails
Courier New 10pt; Resumo recuado 0,8 cm; títulos de seção 13pt bold à esquerda;
corpo justificado com 1ª linha recuada 1,27 cm (exceto 1º parágrafo da seção).
Saída: Artigo_ObiOne_SBC.docx
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from pathlib import Path

TIMES = "Times New Roman"
COURIER = "Courier New"
HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "Artigo_ObiOne_SBC.docx")
PRINTS = Path(HERE).parent / "apresentacoes" / "prints"

doc = Document()

# ---- Página: A4 + margens SBC ----
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21.0)
sec.top_margin = Cm(3.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(3.0)
sec.right_margin = Cm(3.0)

normal = doc.styles["Normal"]
normal.font.name = TIMES
normal.font.size = Pt(12)


def _run(p, text, *, font=TIMES, size=12, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return r


def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    _run(p, text, size=16, bold=True)


def centered(text, *, size=12, bold=False, font=TIMES, before=0, after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    _run(p, text, size=size, bold=bold, font=font)


def abstract_block(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(p, f"{label} ", bold=True, italic=True)
    _run(p, text, italic=True)


def heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(p, text, size=13, bold=True)


def subheading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(p, text, size=12, bold=True)


def body(text, first=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    if not first:
        p.paragraph_format.first_line_indent = Cm(1.27)
    _run(p, text)


def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    _run(p, text, font="Helvetica", size=10, bold=True)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    if widths:
        t.autofit = False
        t.allow_autofit = False
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.paragraphs[0].clear()
        _run(c.paragraphs[0], h, size=11, bold=True)
        if widths:
            c.width = Cm(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].clear()
            _run(cells[i].paragraphs[0], val, size=11)
            if widths:
                cells[i].width = Cm(widths[i])


def reference(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    _run(p, text, size=12)


def add_hyperlink(paragraph, url, text):
    """Insere um hyperlink clicável (azul, sublinhado) num parágrafo."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hlink = OxmlElement("w:hyperlink")
    hlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "1155CC"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    rFonts = OxmlElement("w:rFonts"); rFonts.set(qn("w:ascii"), "Times New Roman"); rFonts.set(qn("w:hAnsi"), "Times New Roman"); rPr.append(rFonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "24"); rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; new_run.append(t)
    hlink.append(new_run)
    paragraph._p.append(hlink)


def link_line(label, url):
    """Parágrafo 'label: <hyperlink>' no corpo (Times 12, justificado)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(3)
    _run(p, f"{label}: ")
    add_hyperlink(p, url, url)


def figure(img_path, caption_text, width_cm=13.0):
    """Insere imagem centralizada + legenda estilo SBC (Helvetica 10 bold)."""
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(); run.add_picture(str(img_path), width=Cm(width_cm))
    caption(caption_text)


def flow_box(text):
    """Caixa única (bordada) com o fluxo em uma linha, setas → entre etapas.
    Um 'diagrama simples' de pipeline, sem imagem."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.autofit = False
    t.allow_autofit = False
    cell = t.rows[0].cells[0]
    cell.width = Cm(15.0)
    cell.paragraphs[0].clear()
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(cell.paragraphs[0], text, size=11)


# ====================== CONTEÚDO ======================
title("ObiOne: um observatório de projetos viabilizado por IA Generativa")
centered("Bruno Rocha, Cynthia Oliveira, Moisés Júnior, Raniel Silva", size=12, bold=True, after=12)
centered("Escola Politécnica de Pernambuco — Universidade de Pernambuco (UPE)", after=0)
centered("Recife — PE — Brasil", after=0)
centered("Orientador: Prof. Ivaldir Honório de Farias Júnior", after=6)
centered("{bruno.rocha, cynthia.oliveira, moises.junior, raniel.silva}@upe.br",
         font=COURIER, size=10, before=6, after=6)

abstract_block(
    "Resumo.",
    "Consultorias acumulam conhecimento valioso a cada projeto, mas raramente dispõem "
    "de um sistema leve o bastante para preservá-lo e reaproveitá-lo entre clientes. "
    "Este artigo apresenta o ObiOne, um observatório de projetos construído "
    "sobre o Modelo de Observatório de Projetos (MPO) e potencializado por IA "
    "Generativa. O trabalho segue a Design Science Research: além de construir o "
    "artefato, avalia seu funcionamento real (ciclo exercitado de ponta a ponta com IA "
    "da OpenAI) e a percepção de valor em duas rodadas de piloto, com oito consultores "
    "no total. No acumulado os resultados foram positivos (média 4,3 de 5); a segunda "
    "rodada, após a inclusão de um onboarding e com um público mais diverso, foi mais "
    "crítica (4,1 ante 4,48) e manteve a clareza inicial (3,8) como o principal ponto de "
    "atenção. A principal contribuição é a demonstração empírica de que o MPO é "
    "implementável com IA generativa em uma consultoria real.")

# 1. Introdução
heading("1. Introdução")
body("Consultorias entregam, ano após ano, dezenas de projetos para clientes distintos, "
     "e cada projeto produz um conhecimento valioso: o que funcionou, os riscos que se "
     "materializaram e as decisões que mudaram o rumo do trabalho. Esse conhecimento, "
     "porém, tende a permanecer tácito, preso a quem viveu o projeto, e raramente é "
     "capturado de forma reaproveitável. Revisões sistemáticas sobre gestão do "
     "conhecimento mostram que a captura, a análise e a aplicação de lições aprendidas "
     "dependem mais de fatores culturais e organizacionais do que das ferramentas "
     "adotadas (Henz, 2024; Kamudyariwa et al., 2025). O obstáculo prático é o custo: "
     "manter um repositório vivo de conhecimento exige um esforço contínuo que poucas "
     "consultorias de pequeno e médio porte conseguem sustentar. A consequência é "
     "conhecida: erros se repetem entre projetos, há retrabalho e parte do aprendizado se "
     "perde na transição de um caso para o outro.", first=True)
body("Observatórios de projetos são uma resposta a esse problema. São sistemas de "
     "informação que apoiam a coleta, a organização, o armazenamento, a análise e a "
     "publicação de observações, sistematizando a transparência sobre o andamento e as "
     "decisões dos projetos (Vieira et al., 2021). O Modelo de Observatório de Projetos "
     "(MPO) consolida essa abordagem em um conjunto de conceitos hierárquicos que orienta "
     "a concepção desses sistemas (Vieira, 2022; de Farias Junior et al., 2025). O que "
     "distingue um observatório de uma ferramenta comum de gestão é justamente tratar a "
     "observação, e não apenas a execução, como objeto de primeira classe: ferramentas "
     "usuais, como quadros de tarefas e painéis, registram o que foi feito, mas não "
     "capturam o porquê das decisões nem transformam observações em conhecimento "
     "compartilhado.")
body("No contexto de uma consultoria, o portfólio se organiza por áreas de atuação e os "
     "clientes se renovam a cada engajamento, o que torna o reaproveitamento de "
     "aprendizados entre casos especialmente valioso. Além disso, o cliente não é apenas "
     "destinatário do resultado: ele participa das decisões e detém informação que "
     "enriquece a observação. Abrir o observatório, de forma controlada, à participação "
     "do cliente cria uma dimensão comunitária que a literatura de observatórios, "
     "centrada na organização executora, ainda não explora.")
body("Diante desse cenário, sobressaem-se duas lacunas. A primeira é técnica e empírica: "
     "embora o MPO tenha sido validado conceitualmente e em estudos de caso (de Farias "
     "Junior et al., 2025), nenhuma implementação conhecida o operacionaliza com IA "
     "generativa. A segunda é comunitária: os observatórios descritos na literatura "
     "tratam da organização executora, sem explorar a participação do cliente como ator "
     "do ciclo de conhecimento. A IA generativa surge, nesse ponto, como uma janela de "
     "oportunidade, pois reduz o custo de extrair e sintetizar informação textual, tarefa "
     "antes cara e manual.")
body("Este artigo investiga como a IA generativa pode viabilizar um observatório de "
     "projetos, de modo a reduzir a fricção de manutenção e promover o "
     "engajamento entre a organização executora e seus clientes. Desse objetivo geral "
     "derivam três objetivos específicos: demonstrar a viabilidade técnica de "
     "operacionalizar o MPO com IA generativa; construir e exercitar a dimensão "
     "comunitária, em que consultoria e clientes debatem observações e consolidam "
     "aprendizados; e avaliar a percepção de valor da solução por usuários reais.")
body("Para responder a essa pergunta, o trabalho apresenta o ObiOne, um observatório "
     "de projetos construído sobre o MPO e potencializado por IA generativa, e segue a "
     "Design Science Research, que trata a construção do artefato como forma legítima de "
     "investigação. O restante do artigo está organizado como segue: a Seção 2 apresenta "
     "a fundamentação teórica; a Seção 3, o método; a Seção 4, a implementação do "
     "artefato; a Seção 5, os resultados da avaliação; a Seção 6, a discussão e as lições "
     "aprendidas; e a Seção 7, a conclusão.")

# 2. Fundamentação Teórica
heading("2. Fundamentação Teórica")
subheading("2.1. Observatórios de projetos e o MPO")
body("A observação é um elemento antigo da gestão, e a transparência é o construto que, "
     "na literatura recente, traduz a observação no contexto de projetos: tornar visíveis "
     "o andamento, as decisões e os riscos fortalece o alinhamento entre a organização e "
     "seus interessados (Vieira et al., 2021). Observatórios de projetos são os sistemas "
     "de informação que sistematizam essa transparência, apoiando a coleta, a "
     "organização, o armazenamento, a análise e a publicação de observações.", first=True)
body("O Modelo de Observatório de Projetos (MPO) é um modelo conceitual para esses "
     "sistemas, organizado a partir de 61 conceitos estruturados em três níveis "
     "hierárquicos, do geral ao específico (de Farias Junior et al., 2025). Sua versão de "
     "tese sistematiza os atributos de observação no Quadro 37, que reúne 44 atributos "
     "distribuídos em oito dimensões, abrangendo desde dados estruturais do projeto até "
     "registros narrativos como escopo, riscos e lições aprendidas (Vieira, 2022). O "
     "modelo parte de uma perspectiva sociotécnica e relaciona transparência e governança "
     "a sistemas de apoio à gestão de projetos; foi avaliado conceitualmente e em estudos "
     "de caso (de Farias Junior et al., 2025), mas ainda não havia sido operacionalizado "
     "com apoio de IA generativa.")
subheading("2.2. Gestão do conhecimento em projetos")
body("A gestão do conhecimento trata da captura e do reaproveitamento do que se aprende "
     "ao longo do trabalho, e as lições aprendidas são uma de suas práticas centrais em "
     "projetos. Revisões sistemáticas recentes indicam que a implementação efetiva "
     "depende sobretudo de fatores humanos e organizacionais, como cultura, apoio "
     "gerencial e melhoria contínua, e que, sem esses fatores, as práticas de lições "
     "aprendidas tendem a falhar independentemente das ferramentas empregadas (Henz, "
     "2024).", first=True)
body("Em projetos complexos, a aprendizagem organizacional se sustenta quando a captura, "
     "a análise e a aplicação de lições são conduzidas de forma sistemática, e não "
     "pontual (Kamudyariwa et al., 2025). Esse é o ponto em que as ferramentas usuais de "
     "gestão mostram seu limite: elas registram o que foi feito, mas não o porquê, e não "
     "oferecem um lugar nem um ritual para que o conhecimento seja debatido e "
     "consolidado. O observatório atua exatamente sobre essa lacuna.")
subheading("2.3. IA Generativa como assistente")
body("Modelos de linguagem de grande porte têm sido estudados como apoio a tarefas de "
     "análise textual, entre elas a codificação e a síntese qualitativa. Estudos recentes "
     "mostram que a IA acelera a identificação de temas descritivos e reduz o esforço "
     "operacional, mas perde nuances que dependem de conhecimento contextual humano (CHI, "
     "2026; medRxiv, 2024).", first=True)
body("Daí emerge um consenso de parceria guiada: a IA assume tarefas estruturadas "
     "enquanto o humano permanece como líder intelectual da interpretação, princípio "
     "conhecido como human-in-the-loop. Para um observatório, isso sugere um papel "
     "assistivo para a IA, capaz de reduzir a fricção de registrar e sintetizar "
     "conhecimento sem substituir o julgamento de quem conduz o projeto. É essa a posição "
     "adotada pelo ObiOne.")

# 3. Método
heading("3. Método")
subheading("3.1. Design Science Research")
body("A construção do ObiOne segue a Design Science Research, que estabelece o artefato "
     "como forma legítima de pesquisa (Hevner et al., 2004) e organiza o trabalho em "
     "atividades de identificação do problema, definição de objetivos, design e "
     "desenvolvimento, demonstração, avaliação e comunicação (Peffers et al., 2007). Por "
     "se tratar de uma pergunta de viabilidade, responder exige construir o sistema, "
     "colocá-lo em uso e observar o resultado, e não apenas coletar opiniões sobre uma "
     "descrição.", first=True)
subheading("3.2. Desenho da avaliação")
body("A avaliação combinou a demonstração de uso real do artefato com a medição da "
     "percepção de valor. A percepção foi medida por um instrumento de doze afirmações em "
     "escala Likert de 1 a 5 e três perguntas abertas, aplicado em duas rodadas de piloto "
     "com quatro consultores cada, após um walkthrough do sistema. A segunda rodada "
     "ocorreu depois da inclusão de um onboarding de primeiro acesso e reuniu um público "
     "mais diverso. Como as duas rodadas usaram participantes distintos, a comparação é "
     "exploratória e carrega um confundidor: a variação entre rodadas mistura o efeito das "
     "mudanças no produto com a diferença de composição das coortes. A avaliação foi "
     "conduzida na perspectiva do consultor, perfil central do observatório; a aplicação, "
     "porém, oferece telas para os três perfis, consultor, administrador e cliente, "
     "conforme o Apêndice C. Os resultados são reportados como casos, sem inferência "
     "estatística.", first=True)
subheading("3.3. Participantes")
body("As duas rodadas somaram oito participantes de uma mesma consultoria, quatro em cada "
     "rodada. Quanto ao papel, sete eram consultores e um era gestor. As ferramentas de "
     "acompanhamento de projetos usadas no dia a dia eram sobretudo o Trello, por quatro "
     "participantes, e planilhas, por três. A familiaridade com gestão de projetos era, na "
     "maioria, de nível médio. A Tabela 1 resume o perfil.", first=True)
caption("Tabela 1. Perfil dos participantes (N=8, duas rodadas de quatro).")
table(
    ["Dimensão", "Distribuição"],
    [
        ["Rodadas", "4 na primeira; 4 na segunda"],
        ["Papel", "7 consultores; 1 gestor"],
        ["Ferramenta de projetos usada hoje", "4 Trello; 3 planilhas; 1 outra"],
        ["Familiaridade com gestão de projetos", "2 alto; 5 médio; 1 baixo"],
    ],
    widths=[6.5, 8.5])

# 4. Implementação
heading("4. Implementação")
subheading("4.1. Requisitos e rastreabilidade ao MPO")
body("O ObiOne foi desenvolvido para uma consultoria de marketing, que atua como "
     "organização executora e curadora, e seus clientes, que acessam cada um o próprio "
     "projeto. A elicitação e a especificação seguiram práticas usuais de engenharia de "
     "requisitos (Sommerville, 2016), e cada requisito funcional foi ancorado a uma "
     "característica ou processo do MPO, de modo que o artefato implementasse o modelo, e "
     "não apenas se inspirasse nele. A Tabela 2 apresenta uma amostra; a rastreabilidade "
     "completa está no Apêndice A.", first=True)
caption("Tabela 2. Amostra da rastreabilidade requisito → MPO → implementação.")
table(
    ["Requisito (ObiOne)", "Âncora no MPO", "Implementação"],
    [
        ["Catálogo de atributos de observação", "Quadro 37: 44 atributos em 8 dimensões (Vieira, 2022)", "mpo/MpoCatalog"],
        ["Cobertura de observação por projeto", "Característica Abrangência (Vieira, 2022)", "GET /projects/{id}/coverage"],
        ["Governança de acesso por papel", "Característica Segurança (Vieira, 2022)", "filtro por papel (consultor/cliente)"],
        ["Registro e acompanhamento", "Processos Acompanhar e Avaliar (Vieira, 2022)", "ciclo observação → conversa → aprendizado"],
        ["Consolidação de aprendizados", "Transparência e disseminação (de Farias Junior et al., 2025)", "comunidade por domínio"],
    ])
body("A cobertura resultante abrange os 44 atributos de observação previstos pelo MPO, "
     "distribuídos em oito dimensões. "
     "Trata-se de cobertura arquitetural: o sistema provê os campos e os fluxos "
     "correspondentes. A avaliação empírica da qualidade da extração é uma frente "
     "prevista no protocolo, ainda não executada (Seção 5.4).", first=True)
subheading("4.2. Arquitetura")
body("O ObiOne é uma aplicação web dividida em backend e frontend. O backend usa Java 21 "
     "e Spring Boot, com os módulos web, data-jpa, security e validation, e segue uma "
     "organização por contexto de domínio em camadas: controladores REST finos, serviços "
     "com a lógica de negócio, repositórios Spring Data, entidades JPA e objetos de "
     "transferência de dados com mapeadores dedicados. A persistência de desenvolvimento "
     "usa um banco H2 em arquivo, com PostgreSQL previsto para produção. O frontend usa "
     "React com TanStack Start e roteamento baseado em arquivos, construído com Vite; o "
     "estado de servidor é gerido por react-query e os formulários por react-hook-form com "
     "validação por esquema. A API responde sob o caminho base barra-api, e a aplicação "
     "adota um padrão de origem única: o servidor de desenvolvimento serve a interface e "
     "encaminha as chamadas de API ao backend, o que simplifica o acesso remoto para as "
     "sessões de validação.", first=True)
subheading("4.3. Camada de IA")
body("A camada de IA é o principal diferencial do ObiOne e atua de forma assistiva sobre "
     "o ciclo de observação, conversa e aprendizado. Está organizada em cinco papéis; em "
     "todos, a saída é uma sugestão, nunca uma ação publicada automaticamente. A Tabela 3 "
     "resume os papéis, com suas entradas e saídas.", first=True)
caption("Tabela 3. Papéis da camada de IA, com entradas e saídas.")
table(
    ["Papel", "Função", "Entrada", "Saída"],
    [
        ["Categorizadora", "Sugere o domínio do projeto", "resumo, objetivo, domínios disponíveis", "domínio sugerido e confiança"],
        ["Observadora", "Sugere observações ancoradas no MPO", "resumo, objetivo, lente MPO, atributos prioritários, observações já registradas", "observações mapeadas a atributos, com impacto e trecho literal"],
        ["Sintetizadora", "Rascunha um aprendizado a partir da conversa", "título, pergunta, contribuições", "rascunho com resumo, evidência e recomendação"],
        ["Conectora", "Sintetiza padrões entre projetos do domínio (implementada; não avaliada)", "resumos dos projetos do domínio", "padrões e lições anonimizados"],
        ["Configuradora", "Sugere o setup inicial no cadastro", "nome, descrição, objetivo", "domínio, atributos e fenômenos esperados"],
    ],
    widths=[3.0, 4.5, 3.75, 3.75])
body("O processamento de um projeto segue um fluxo comum. A partir do texto do projeto, o "
     "serviço assistente reúne o contexto e o analisa à luz dos atributos de observação "
     "do MPO; aciona a IA, que devolve uma sugestão estruturada; registra essa sugestão "
     "de forma auditável, com sua proveniência; e a devolve ao consultor para revisão. A "
     "Figura 1 ilustra esse fluxo.", first=True)
flow_box("Descrição do projeto  →  Análise pela IA à luz do MPO  →  "
         "Sugestão estruturada e auditável  →  Revisão do consultor  →  "
         "Observação ou aprendizado publicado")
caption("Figura 1. Pipeline da camada de IA.")
body("Quatro técnicas sustentam a confiabilidade das sugestões. A primeira é a saída "
     "estruturada: o modelo é obrigado a responder no formato de um objeto de dados, que "
     "o sistema mapeia diretamente, sem interpretação livre do texto. A segunda é o "
     "grounding pela lente do MPO, reforçado por instruções que orientam o modelo a não "
     "inventar atributos fora da lista fornecida e a citar o trecho literal do resumo que "
     "motivou cada observação. A terceira é uma validação determinística em código: na "
     "configuração inicial de um projeto, identificadores de atributo ou de domínio "
     "inexistentes no catálogo são descartados antes de a resposta ser devolvida. A "
     "quarta é a consciência do estado do observatório: ao sugerir observações, o modelo "
     "recebe a lista do que o projeto já registra e é instruído a não repetir aspectos em "
     "observação, e um filtro determinístico descarta, em código, qualquer sugestão cujo "
     "atributo já esteja coberto, mesmo que o modelo desobedeça. O provedor é "
     "configurável: um modo determinístico, sem chave e "
     "voltado a testes, e o provedor da OpenAI, com o modelo gpt-5.4-mini e temperatura "
     "baixa, para uso real.", first=True)
body("A IA nunca escreve diretamente nas observações ou nos aprendizados. Ela apenas "
     "sugere e registra cada sugestão em um log de auditoria, com o provedor, o modelo, o "
     "instante e a indicação de aceite, o que dá reprodutibilidade ao uso da IA. A "
     "persistência só ocorre quando o consultor aceita a sugestão, e a observação é então "
     "gravada com a origem marcada como assistida pela IA. A taxa de aceite por tipo de "
     "sugestão é observável no sistema, permitindo acompanhar o quanto as sugestões são de "
     "fato aproveitadas. No caso da Conectora, a própria síntese fica persistida nesse "
     "log: o sistema recupera a última versão por domínio, com a data de geração, e "
     "oferece a regeração sob demanda, tratando o resultado como um artefato de "
     "conhecimento versionado, e não como uma saída efêmera.", first=True)
subheading("4.4. Prototipação")
body("Antes do desenvolvimento final, as telas foram prototipadas com apoio de "
     "ferramentas de geração assistida por IA no ecossistema React, incluindo o Lovable, "
     "que produziu o scaffold inicial da interface e da sua configuração de build. A "
     "prototipação rápida corresponde à atividade de design e desenvolvimento da Design "
     "Science Research e serviu de insumo concreto para as validações com os "
     "orientadores, encurtando o ciclo entre uma ideia de tela e uma versão navegável.",
     first=True)
subheading("4.5. Governança por papel")
body("O acesso ao observatório é semi-aberto e governado pelo papel do usuário. As "
     "leituras exigem autenticação; as mutações são restritas aos papéis de consultor e "
     "administrador, enquanto o cliente contribui nas conversas e enxerga apenas o seu "
     "próprio projeto. O consultor conduz a curadoria e vê todo o portfólio; o "
     "administrador acumula as permissões de gestão; o cliente participa da comunidade do "
     "seu caso sem acesso às ações de equipe nem à visão consolidada do portfólio. A "
     "configuração dessas permissões é, ela própria, governada por regras assimétricas: o "
     "perfil do administrador é fixo, com todas as permissões habilitadas e imutáveis; "
     "apenas o administrador altera as permissões do consultor; e o consultor gerencia "
     "apenas as do cliente, com as regras aplicadas no backend, e não apenas na "
     "interface. A integridade da evidência segue o mesmo princípio: uma observação com "
     "conversa vinculada não pode ser excluída, preservando o registro que sustenta o "
     "debate da comunidade. Esse "
     "arranjo garante o isolamento entre clientes e materializa, na prática, o acesso "
     "semi-aberto previsto no MPO. As telas correspondentes a cada perfil estão no "
     "Apêndice C.", first=True)
subheading("4.6. Jornada do usuário e construção do MVP")
body("A jornada central percorre quatro momentos: o cadastro de um projeto assistido por "
     "um wizard com apoio de IA, o registro de observações ancoradas no MPO, a conversa da "
     "comunidade sobre essas observações e a consolidação de aprendizados reaproveitáveis. "
     "Um feed reflete a atividade recente e a cobertura de cada projeto frente ao MPO é "
     "consultável. O MVP foi construído em torno desse fluxo e entregou o cadastro com "
     "extração assistida, a organização da comunidade por domínio, o detalhe do projeto "
     "com os atributos do MPO, a consolidação de aprendizados com apoio da Sintetizadora, "
     "o feed de novidades e, após a primeira rodada de validação, um onboarding de "
     "primeiro acesso que apresenta o objetivo do sistema e o ciclo de conhecimento.",
     first=True)

# 5. Resultados
heading("5. Resultados")
subheading("5.1. Viabilidade técnica")
body("Em junho de 2026, o ciclo foi exercitado de ponta a ponta com a IA da OpenAI e "
     "dados de simulação. As quatro frentes assistivas operaram de forma integrada: "
     "cadastro, sugestão de observações priorizadas pelos riscos declarados, aceite pelo "
     "consultor, alimentação da cobertura, conversa com participação do cliente e "
     "atualização do feed. A governança por papel foi confirmada na interface.", first=True)
subheading("5.2. Percepção de valor")
body("A percepção de valor foi medida em duas rodadas de piloto com quatro consultores "
     "cada, aplicando o mesmo instrumento de doze afirmações Likert e três perguntas "
     "abertas. A primeira rodada ocorreu após um walkthrough curto; a segunda, após a "
     "inclusão de um onboarding de primeiro acesso e outros ajustes, com um público mais "
     "diverso (incluindo um gestor e mais usuários de planilhas). A primeira rodada "
     "obteve média 4,48 de 5, com 44 de 48 respostas nas notas 4 ou 5 e 29 máximas. A "
     "segunda foi mais crítica: média 4,1, com 36 de 48 respostas positivas e 19 "
     "máximas. No acumulado dos oito participantes, a média foi 4,3 de 5, com 80 de 96 "
     "respostas positivas. A Tabela 4 apresenta o comparativo por dimensão entre as duas "
     "rodadas.", first=True)
caption("Tabela 4. Médias por dimensão nas duas rodadas (escala 1 a 5, N=4 por rodada).")
table(
    ["Dimensão", "1ª rodada", "2ª rodada", "Δ"],
    [
        ["Clareza", "3,8", "3,8", "0,0"],
        ["Organização", "4,8", "4,0", "-0,8"],
        ["Usabilidade", "4,5", "4,5", "0,0"],
        ["Conteúdo", "4,8", "4,8", "0,0"],
        ["Diferenciação", "4,2", "3,8", "-0,5"],
        ["Ciclo de conhecimento", "4,0", "4,2", "+0,2"],
        ["Comunidade", "4,8", "4,3", "-0,5"],
        ["Aprendizados", "5,0", "4,2", "-0,8"],
        ["IA assistiva", "4,5", "4,0", "-0,5"],
        ["Portfólio", "4,2", "4,2", "0,0"],
        ["Governança", "5,0", "3,8", "-1,2"],
        ["Intenção de uso", "4,2", "4,0", "-0,2"],
    ],
    widths=[6.0, 3.0, 3.0, 3.0])
body("A segunda rodada leu o produto de forma mais crítica em governança (-1,2), "
     "organização e aprendizados (-0,8 cada), comunidade, diferenciação e IA assistiva "
     "(-0,5 cada). Conteúdo, usabilidade, portfólio e, notadamente, a clareza "
     "permaneceram estáveis, e o ciclo de conhecimento evoluiu (+0,2). A queda é "
     "atribuída, em parte, a um público mais diverso e crítico e, em parte, à fricção "
     "inicial que persiste: isolar o efeito do onboarding exigiria um desenho controlado "
     "com os mesmos usuários antes e depois. Ainda assim, o valor central se manteve nas "
     "respostas abertas, em que a IA passou a aparecer como força ("
     "“aprendizado consolidado através da IA”) e o produto foi descrito como "
     "amigável e que integra tecnologia e educação. O feedback crítico, por sua vez, "
     "concentrou-se na experiência inicial: um participante relatou que “ficou um pouco "
     "confusa a explicação sobre algumas funcionalidades”, e os pedidos de melhoria "
     "convergiram para “navegação mais intuitiva, instruções mais claras e interface mais "
     "simples”.", first=True)
subheading("5.3. Benefícios esperados")
body("Os benefícios esperados com a solução incluem a redução do custo de manutenção do "
     "conhecimento entre projetos, o fortalecimento do relacionamento com os clientes por "
     "meio de maior engajamento e comunicação na comunidade, e o reaproveitamento de "
     "aprendizados já consolidados em novos projetos do mesmo domínio.", first=True)
subheading("5.4. Limitações")
body("As duas rodadas somam oito participantes de uma mesma consultoria, em ciclos "
     "distintos e com composição diferente; os resultados são reportados como casos, "
     "sem inferência estatística, e a comparação entre rodadas é exploratória. A clareza "
     "permaneceu como o principal ponto de atenção (3,8 nas duas rodadas): a inclusão de "
     "um onboarding de primeiro acesso não moveu esse indicador, o que sugere que a "
     "orientação inicial precisa ir além de um passo introdutório, na direção de uma "
     "navegação mais guiada. Como a segunda rodada usou participantes distintos, não é "
     "possível separar o efeito do onboarding da mudança de público. Por fim, a avaliação "
     "da qualidade da extração dos atributos do MPO, prevista no protocolo, ainda não foi "
     "executada.", first=True)

# 6. Discussões e Lições Aprendidas
heading("6. Discussões e Lições Aprendidas")
body("Os resultados confirmam o MPO como base válida e mostram que ele é implementável "
     "com IA generativa a um custo viável para uma consultoria de pequeno e médio porte. "
     "Estudos anteriores avaliaram o modelo conceitualmente e em casos (de Farias Junior "
     "et al., 2025); o ObiOne acrescenta uma implementação operacional com IA. A IA "
     "reduz a fricção de iniciar o ciclo; à objeção de que criaria dependência, o desenho "
     "responde com o human-in-the-loop, e as respostas abertas sustentam que o valor "
     "percebido está no que a comunidade produz e no aprendizado consolidado com apoio da "
     "IA. A segunda rodada, mais crítica, deixa claro que a restrição dominante para a "
     "adoção não está no ciclo em si, mas na experiência inicial: enquanto a clareza e a "
     "navegação não forem resolvidas, o valor demora a ser percebido. O fato de o "
     "onboarding não ter movido a clareza indica que o próximo passo é uma navegação "
     "guiada, com menu evidente, fluxo em etapas e exemplos práticos.", first=True)
body("Essa leitura já produziu um terceiro ciclo de design, posterior à segunda rodada. "
     "O cadastro tornou-se IA-first, com a descrição textual como único insumo e "
     "degradação graciosa quando a IA não responde; o ciclo de observação, conversa e "
     "consolidação passou a acontecer em uma única tela, com um clique para abrir a "
     "conversa e o diálogo de consolidação reduzido ao essencial; e as telas foram "
     "enxugadas para diminuir a carga de leitura. Essas mudanças respondem diretamente à "
     "restrição de clareza apontada pela avaliação, mas ainda não foram reavaliadas: a "
     "rodada controlada prevista nas questões em aberto é o teste natural desse "
     "redesenho.")
body("Como decisões arquiteturais, destacam-se a IA estritamente assistiva e a "
     "governança por papel, que viabiliza o acesso semi-aberto com isolamento entre "
     "clientes. Entre vantagens e limitações das ferramentas, o desenvolvimento em código "
     "deu controle sobre o pipeline e a governança, ao custo de mais esforço do que uma "
     "abordagem low-code. A experiência de uso do MPO mostrou que traduzir 44 atributos "
     "em uma interface sem jargão é mais difícil do que implementá-los.")
body("A lição mais marcante foi de escopo. O projeto passou por um pivô: a proposta "
     "inicial foi reescopada para refinar o propósito da solução, deslocando o foco de um "
     "extrator de atributos para um observatório de projetos com participação dos "
     "clientes. Esse refino exigiu bastante "
     "trabalho ao longo de várias validações com os orientadores, e foi ele, mais do que "
     "qualquer ganho de ferramenta, que destravou o valor percebido. A IA generativa "
     "acelerou a construção, mas mostrou um limite claro: sem uma definição nítida do que "
     "se está construindo, a velocidade da IA não leva a lugar nenhum; ela amplifica a "
     "direção que já existe, não a substitui.")
body("Amarradas a essa jornada, as lições de equipe se somam às de escopo. No plano operacional, o escopo amplo frente ao "
     "prazo exigiu disciplina de cronograma e trabalho de preparação de dados e "
     "integração contínua, com deslizes de data corrigidos ao longo do projeto. No plano "
     "técnico, os principais desafios foram a integração da IA, a configuração de acesso "
     "remoto para validação e a manutenção do ambiente. No plano conceitual, entender o "
     "MPO em profundidade foi determinante. No plano de equipe, coordenar quatro "
     "integrantes e integrar frentes de extração, frontend, avaliação e escrita exigiu "
     "comunicação constante.")

# 7. Conclusão
heading("7. Conclusão")
body("O custo de manter um observatório de projetos caiu com a IA generativa, mas o "
     "diferencial de valor é a comunidade, e isso não é substituível por tecnologia. "
     "Como síntese da experiência, o ObiOne demonstra empiricamente que o MPO é "
     "implementável com IA generativa em uma consultoria real, e que a participação de "
     "consultoria e clientes em um ciclo comum de conhecimento distingue o observatório "
     "de uma ferramenta de gestão. A principal contribuição é essa demonstração "
     "empírica.", first=True)
subheading("Questões em aberto")
body("A leitura comparativa das duas rodadas deixa cinco questões que delimitam os gaps "
     "da pesquisa e orientam sua continuidade: (1) o onboarding melhora a clareza quando "
     "medido nos mesmos usuários antes e depois, isolando o efeito da mudança de público? "
     "(2) uma navegação guiada, com menu evidente, fluxo em etapas e exemplos práticos, "
     "eleva a clareza acima do patamar observado de 3,8? (3) qual a acurácia e a "
     "fidelidade da extração dos atributos do MPO pela IA, conforme o protocolo ainda não "
     "executado? (4) a IA melhora a qualidade dos aprendizados consolidados, e não apenas "
     "reduz a fricção de iniciar o ciclo? (5) o padrão de valor percebido se mantém em "
     "outros domínios, com amostra maior e uso prolongado, e o que explica a queda na "
     "percepção de governança na segunda rodada?", first=True)
body("Essas questões orientam os trabalhos futuros: realizar uma nova rodada controlada, "
     "com os mesmos usuários antes e depois do redesenho IA-first e da navegação guiada; "
     "ampliar a validação com mais participantes e domínios; executar o protocolo de "
     "avaliação da extração do MPO; e avaliar a síntese cross-projeto (Conectora), já "
     "implementada com anonimização e com cada síntese persistida no log de auditoria, "
     "cuja avaliação de valor permanece em aberto.")

# Referências
heading("Referências")
refs = [
    'de Farias Junior, I. H., Vieira, J. K. M., de Moura, H. P. and Sampaio, L. T. (2025) "A Conceptual Model for Project Observatories", IEEE Access, v. 13. DOI: 10.1109/ACCESS.2025.3589743.',
    'Henz, P. (2024) "Knowledge management implementation: A systematic literature review", Knowledge and Process Management. DOI: 10.1002/kpm.1780.',
    'Hevner, A. R., March, S. T., Park, J. and Ram, S. (2004) "Design Science in Information Systems Research", MIS Quarterly, v. 28, n. 1, p. 75-106. DOI: 10.2307/25148625.',
    'Kamudyariwa, X. B., Osobajo, O. A., Oke, A. and Adebayo, Y. (2025) "Application of the systemic lessons learned knowledge model to learning in complex projects". DOI: 10.1177/13505076251339433.',
    'Peffers, K., Tuunanen, T., Rothenberger, M. A. and Chatterjee, S. (2007) "A Design Science Research Methodology for Information Systems Research", Journal of Management Information Systems, v. 24, n. 3, p. 45-77. DOI: 10.2753/MIS0742-1222240302.',
    'Sommerville, I. (2016) "Software Engineering", 10ª ed., Pearson Education, Boston.',
    '"Generative AI for Thematic Analysis in a Maternal Health Study: Coding Semi-structured Interviews using Large Language Models" (2024), medRxiv (preprint). DOI: 10.1101/2024.09.16.24313707.',
    '"Qualitative Coding Analysis through Open-Source Large Language Models: A User Study and Design Recommendations" (2026), In: CHI Conference on Human Factors in Computing Systems, Extended Abstracts. DOI: 10.1145/3772363.3798320.',
    'Vieira, J. K. M. (2022) "Observatórios de Projetos: Um Modelo Conceitual", Tese de Doutorado, Centro de Informática, Universidade Federal de Pernambuco, Recife.',
    'Vieira, J. K. M., de Farias Junior, I. H. and de Moura, H. P. (2021) "Observatories as Transparency Instruments for Projects", In: 16ª Conferência Ibérica de Sistemas e Tecnologias de Informação (CISTI).',
]
for r in refs:
    reference(r)

# Apêndices
heading("Apêndices")
subheading("Apêndice A - Requisitos e rastreabilidade ao MPO")
body("A Tabela A.1 relaciona os principais requisitos funcionais do ObiOne às "
     "características e processos do MPO (Vieira, 2022) e à sua materialização no sistema. "
     "A especificação completa dos requisitos está no repositório.", first=True)
caption("Tabela A.1. Rastreabilidade dos requisitos funcionais ao MPO.")
table(
    ["RF", "Requisito", "Âncora no MPO (Vieira, 2022)", "Implementação"],
    [
        ["RF01", "Autenticar usuário", "Segurança (p. 192)", "auth por token JWT stateless; SecurityConfig"],
        ["RF02", "Perfis e acesso semi-aberto", "Acesso semi-aberto (p. 189); agentes (pp. 200-201)", "papéis consultor/admin/cliente"],
        ["RF03", "Cadastrar projeto (descrição textual)", "Coletar (p. 195)", "wizard de cadastro"],
        ["RF04", "Governança de visibilidade por papel", "Acesso semi-aberto (p. 189); Segurança (p. 192)", "enforcement por papel"],
        ["RF05", "Extrair atributos do MPO via IA", "Quadro 37 (p. 264); Transformar (p. 196)", "camada de IA (Observadora)"],
        ["RF06", "Persistir extração estruturada", "Armazenar (p. 196)", "entidades JPA"],
        ["RF07", "Portfólio perfil-aware", "Abrangência (p. 189); Disponibilizar (p. 196)", "listagem por papel"],
        ["RF08", "Detalhe do projeto", "Projetos (p. 186); Disponibilizar (p. 196)", "detalhe com atributos MPO"],
        ["RF09", "Cobertura do MPO", "Abrangência (p. 189); Avaliar (p. 198)", "GET /projects/{id}/coverage"],
        ["RF10", "Comentar/conversar no projeto", "Interatividade (p. 191); Interagir (p. 198)", "discussões da comunidade"],
        ["RF11", "Feed de novidades", "Acompanhar (p. 198)", "feed temporal in-app"],
        ["RF12", "Consolidar aprendizados (IA)", "Transformar; Categorizar (pp. 196-197)", "camada de IA (Sintetizadora)"],
        ["RF13", "Categorizar por temática/domínio", "Categorizar/Classificar (pp. 196-197)", "camada de IA (Configuradora)"],
    ])
link_line("Repositório", "https://github.com/raniel90/obione")
link_line("Especificação de requisitos", "https://github.com/raniel90/obione/blob/main/atividades/requisitos.md")

subheading("Apêndice B - Arquitetura do Observatório")
body("A arquitetura em camadas do backend, com controladores, serviços, repositórios, "
     "entidades e mapeadores, e o pipeline da camada de IA estão descritos em detalhe nos "
     "documentos de arquitetura do repositório, listados a seguir.", first=True)
link_line("Arquitetura do backend", "https://github.com/raniel90/obione/blob/main/atividades/arquitetura_backend.md")
link_line("Pipeline da camada de IA", "https://github.com/raniel90/obione/blob/main/atividades/arquitetura_pipeline.md")
link_line("Diagrama da arquitetura", "https://github.com/raniel90/obione/blob/main/atividades/arquitetura_diagrama.md")

subheading("Apêndice C - Telas por perfil")
body("A Tabela C.1 resume quais perfis acessam cada tela; em seguida, as Figuras C.1 a "
     "C.6 ilustram as telas principais. O conjunto completo de telas está em "
     "Principais_Telas_ObiOne.pdf, no repositório.", first=True)
caption("Tabela C.1. Telas e perfis que as acessam.")
table(
    ["Tela", "Consultor", "Administrador", "Cliente"],
    [
        ["Observatório (home)", "sim", "sim", "visão limitada"],
        ["Comunidade", "sim", "sim", "sim (seu domínio)"],
        ["Detalhe do projeto", "sim (todos atributos)", "sim", "sim (seu projeto)"],
        ["Feed de novidades", "sim", "sim", "sim"],
        ["Consolidar com IA", "sim", "sim", "não"],
        ["Wizard de cadastro", "sim", "sim", "não"],
    ])
figure(PRINTS / "01-home.png", "Figura C.1. Observatório (home).")
figure(PRINTS / "02-comunidade.png", "Figura C.2. Comunidade.")
figure(PRINTS / "03-detalhe-projeto.png", "Figura C.3. Detalhe do projeto.")
figure(PRINTS / "04-feed.png", "Figura C.4. Feed de novidades.")
figure(PRINTS / "05-consolidar-ia.png", "Figura C.5. Consolidar aprendizado com IA.")
figure(PRINTS / "06-wizard.png", "Figura C.6. Wizard de cadastro de projeto.")
link_line("Telas completas", "https://github.com/raniel90/obione/blob/main/atividades/Principais_Telas_ObiOne.pdf")

doc.save(OUT)
print("OK:", OUT)
