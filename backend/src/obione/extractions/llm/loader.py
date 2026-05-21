"""Minimal .docx → text loader.

Concatenates non-empty paragraphs and flattens table rows into pipe-joined
lines. Good enough for the consultancy `.docx` files in the case study —
no headers/footers, no images, no nested tables expected at runtime.
"""
import io

from docx import Document


def extract_text_from_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
