import io

import pytest
from docx import Document

from obione.extractions.llm.loader import extract_text_from_docx


def _docx_with_paragraphs(paragraphs: list[str]) -> bytes:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.mark.unit
def test_extract_text_concatenates_paragraphs():
    raw = _docx_with_paragraphs(["Linha 1", "Linha 2", "", "Linha 3"])
    text = extract_text_from_docx(raw)
    assert "Linha 1" in text
    assert "Linha 2" in text
    assert "Linha 3" in text


@pytest.mark.unit
def test_extract_text_skips_blank_paragraphs():
    raw = _docx_with_paragraphs(["A", "   ", "B"])
    text = extract_text_from_docx(raw)
    lines = [line for line in text.split("\n") if line]
    assert lines == ["A", "B"]


@pytest.mark.unit
def test_extract_text_includes_table_cells():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Cabecalho"
    table.rows[0].cells[1].text = "Valor"
    table.rows[1].cells[0].text = "Linha"
    table.rows[1].cells[1].text = "Dado"
    buf = io.BytesIO()
    doc.save(buf)
    text = extract_text_from_docx(buf.getvalue())
    assert "Cabecalho" in text
    assert "Valor" in text
    assert "Linha" in text
    assert "Dado" in text


@pytest.mark.unit
def test_extract_text_handles_empty_doc():
    doc = Document()
    buf = io.BytesIO()
    doc.save(buf)
    assert extract_text_from_docx(buf.getvalue()) == ""


@pytest.mark.unit
def test_extract_text_preserves_pt_br_chars():
    raw = _docx_with_paragraphs(["Valença é uma cidade. Açaí também."])
    text = extract_text_from_docx(raw)
    assert "Valença" in text
    assert "Açaí" in text
